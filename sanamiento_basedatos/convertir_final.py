from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

md_path = '/home/michael/sanamiento_basedatos/GHA_Grupo01_MichaelAxel_AguilarPeredo_Justificacion.md'
docx_path = '/home/michael/sanamiento_basedatos/GHA_Grupo01_MichaelAxel_AguilarPeredo_Justificacion.docx'

with open(md_path, 'r') as file:
    content = file.read()

doc = Document()

def add_formatted_text(doc, text, bold=False, italic=False, size=12):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    return para

lines = content.split('\n')

for line in lines:
    stripped = line.strip()
    
    if not stripped:
        doc.add_paragraph()
        continue
    
    if stripped.startswith('# ') and stripped.count('#') == 1:
        doc.add_heading(stripped[2:], 0)
    
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], 1)
    
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], 2)
    
    elif stripped.startswith('**') and stripped.endswith('**') and stripped[2:-2]:
        add_formatted_text(doc, stripped[2:-2], bold=True)
    
    elif stripped.startswith('- '):
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(stripped[2:])
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    
    elif stripped.startswith('  - '):
        para = doc.add_paragraph(style='List Bullet 2')
        run = para.add_run(stripped[4:])
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    
    elif stripped == '---':
        doc.add_paragraph('─' * 50)
    
    elif stripped.startswith('**') and ':' in stripped:
        parts = stripped[2:].split(':', 1)
        if len(parts) == 2:
            para = doc.add_paragraph()
            run1 = para.add_run(parts[0] + ': ')
            run1.bold = True
            run1.font.name = 'Arial'
            run1.font.size = Pt(11)
            run2 = para.add_run(parts[1].strip())
            run2.font.name = 'Arial'
            run2.font.size = Pt(11)
    
    else:
        add_formatted_text(doc, stripped, size=11)

doc.save(docx_path)
print(f'Archivo Word creado correctamente: {docx_path}')
