import markdown2 as md
from docx import Document

# Leyendo el archivo markdown
with open('/home/michael/sanamiento_basedatos/GHA_Grupo01_MichaelAxel_AguilarPeredo_Justificacion.md', 'r') as file:
    md_text = file.read()

# Convirtiendo el markdown a HTML
html_text = md.markdown(md_text)

# Creando un nuevo documento Word
doc = Document()

# Estableciendo estilo sencillo y profesional
for para in doc.paragraphs:
    para.style.font.name = 'Arial'
    para.style.font.size = 12
    para.style.font.bold = False
    para.style.font.italic = False

# Añadiendo el contenido HTML al documento Word
doc.add_paragraph(html_text)

# Guardando el archivo Word
output_path = '/home/michael/sanamiento_basedatos/GHA_Grupo01_MichaelAxel_AguilarPeredo_Justificacion.docx'
doc.save(output_path)

print(f"Archivo guardado en {output_path}")
