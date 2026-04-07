from docx import Document

md_path = '/home/michael/sanamiento_basedatos/GHA_Grupo01_MichaelAxel_AguilarPeredo_Justificacion.md'
docx_path = '/home/michael/sanamiento_basedatos/GHA_Grupo01_MichaelAxel_AguilarPeredo_Justificacion.docx'

with open(md_path, 'r') as file:
    lines = file.readlines()

doc = Document()

for line in lines:
    line = line.rstrip()
    
    if line.startswith('# '):
        doc.add_heading(line[2:], 0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], 1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], 2)
    elif line.startswith('**') and line.endswith('**'):
        doc.add_paragraph(line, style='Intense Quote')
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('  - '):
        doc.add_paragraph(line[4:], style='List Bullet 2')
    elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
        doc.add_paragraph(line[3:], style='List Number')
    elif line == '---':
        doc.add_paragraph('─' * 50)
    elif line.strip() == '':
        continue
    else:
        doc.add_paragraph(line)

doc.save(docx_path)
print(f'Archivo Word creado: {docx_path}')
