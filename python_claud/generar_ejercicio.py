from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='999999'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, italic=False,
              font_name='Arial', size=10, color='000000', align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_normal(doc, text, italic=False, color='222222'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    return p

# ── Tabla UML de una clase ────────────────────────────────────────────────────
# header_hex: color de cabecera, attrs: lista de strings, methods: lista de strings
def make_class_table(doc, class_name, header_hex, attrs, methods, col_width_cm=7):
    rows_count = 1 + 1 + len(attrs) + 1 + len(methods)
    table = doc.add_table(rows=rows_count, cols=1)
    table.style = 'Table Grid'

    w = Cm(col_width_cm)
    for row in table.rows:
        row.cells[0].width = w

    idx = 0

    # Nombre de clase
    cell = table.rows[idx].cells[0]
    cell_text(cell, class_name, bold=True, size=11, color='FFFFFF',
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(cell, header_hex)
    set_cell_borders(cell, header_hex)
    idx += 1

    # Sección Atributos
    cell = table.rows[idx].cells[0]
    cell_text(cell, 'Atributos', bold=True, size=10, color='1F3864')
    set_cell_bg(cell, 'D6E4F0')
    set_cell_borders(cell)
    idx += 1

    for a in attrs:
        cell = table.rows[idx].cells[0]
        cell_text(cell, a, font_name='Courier New', size=9.5)
        set_cell_bg(cell, 'F5F5F5')
        set_cell_borders(cell)
        idx += 1

    # Sección Métodos
    cell = table.rows[idx].cells[0]
    cell_text(cell, 'Métodos', bold=True, size=10, color='1F3864')
    set_cell_bg(cell, 'D6E4F0')
    set_cell_borders(cell)
    idx += 1

    for m in methods:
        cell = table.rows[idx].cells[0]
        cell_text(cell, m, font_name='Courier New', size=9.5)
        set_cell_bg(cell, 'FFFFFF')
        set_cell_borders(cell)
        idx += 1

    return table

def add_gap(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)

# ══════════════════════════════════════════════════════════════════════════════
#  CONTENIDO DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

# Título
add_heading(doc, 'Diagramas de Clase 2 — Solución', level=1)
add_normal(doc,
    'A partir de las figuras del enunciado (círculo gris, rectángulo naranja, '
    'cuadrado azul y elipse amarilla) se definen las siguientes clases, '
    'atributos y jerarquía de herencia usando especialización/generalización '
    '(herencia y polimorfismo).')
add_gap(doc)

# ── 1. Clase base ─────────────────────────────────────────────────────────────
add_heading(doc, '1. Clase Base Abstracta: Figura', level=2)
make_class_table(doc, 'Figura  «abstracta»', '1F3864',
    attrs=[
        '# color        : String',
        '# posicionX    : double',
        '# posicionY    : double',
    ],
    methods=[
        '+ getColor()                   : String',
        '+ setColor(c : String)         : void',
        '+ getArea()                    : double   {abstracto}',
        '+ getPerimetro()               : double   {abstracto}',
        '+ dibujar()                    : void     {abstracto}',
        '+ mover(x : double, y : double): void',
    ],
    col_width_cm=10
)
add_gap(doc)

# ── 2. Subclases ──────────────────────────────────────────────────────────────
add_heading(doc, '2. Subclases Concretas', level=2)
add_normal(doc, 'Cada figura se modela como una subclase que hereda de Figura '
               'e implementa sus métodos abstractos:')
add_gap(doc, 4)

clases = [
    ('Circulo',    '555555',
     ['- radio : double'],
     ['+ getArea() : double', '+ getPerimetro() : double',
      '+ dibujar() : void',   '+ getRadio() : double',
      '+ setRadio(r : double) : void']),

    ('Rectangulo', 'C45911',
     ['- ancho : double', '- alto : double'],
     ['+ getArea() : double', '+ getPerimetro() : double',
      '+ dibujar() : void',   '+ getAncho() : double',
      '+ getAlto() : double']),

    ('Cuadrado',   '2E75B6',
     ['- lado : double'],
     ['+ getArea() : double', '+ getPerimetro() : double',
      '+ dibujar() : void',   '+ getLado() : double',
      '+ setLado(l : double) : void']),

    ('Elipse',     '888800',
     ['- semiEjeA : double', '- semiEjeB : double'],
     ['+ getArea() : double', '+ getPerimetro() : double',
      '+ dibujar() : void',   '+ getSemiEjeA() : double',
      '+ getSemiEjeB() : double']),
]

for nombre, color, attrs, methods in clases:
    add_normal(doc, nombre, italic=False, color=color.ljust(6,'0'))
    make_class_table(doc, nombre, color, attrs, methods, col_width_cm=9)
    add_gap(doc, 8)

# ── 3. Jerarquía ──────────────────────────────────────────────────────────────
add_heading(doc, '3. Jerarquía de Herencia', level=2)

lines = [
    ('Figura  «superclase abstracta»',          '1F3864', True),
    ('  |-- FiguraPlana  «clase intermedia»',   '333333', False),
    ('  |     |-- Circulo',                     '555555', False),
    ('  |     |-- Elipse',                      '777700', False),
    ('  |     `-- Poligono  «clase intermedia»','333333', False),
    ('  |           |-- Rectangulo',            'C45911', False),
    ('  |           `-- Cuadrado  (extiende Rectangulo)', '2E75B6', False),
]

table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.rows[0].cells[0]
cell.width = Cm(15)
set_cell_bg(cell, 'EEF2FA')
set_cell_borders(cell, '2E75B6')
cell.text = ''

for text, color, bold in lines:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(10)
    run.bold       = bold
    run.font.color.rgb = RGBColor.from_string(color)

add_gap(doc)

# ── 4. Polimorfismo ───────────────────────────────────────────────────────────
add_heading(doc, '4. Polimorfismo', level=2)
add_normal(doc, 'Los métodos abstractos getArea(), getPerimetro() y dibujar() '
               'son implementados de forma distinta por cada subclase:')

add_bullet(doc, 'Circulo.getArea()     = π × radio²')
add_bullet(doc, 'Rectangulo.getArea()  = ancho × alto')
add_bullet(doc, 'Cuadrado.getArea()    = lado²  (donde ancho = alto = lado)')
add_bullet(doc, 'Elipse.getArea()      = π × semiEjeA × semiEjeB')

add_gap(doc, 6)
add_normal(doc,
    'Una variable de tipo Figura puede referenciar cualquier subclase concreta. '
    'Al invocar getArea() o dibujar() se ejecuta el método de la subclase '
    'correspondiente — esto es polimorfismo en tiempo de ejecución.')

add_gap(doc, 4)
p = add_normal(doc,
    'Nota: Cuadrado extiende Rectangulo porque un cuadrado ES un rectángulo '
    'con ancho = alto. Se sobreescriben setAncho() y setAlto() para mantener '
    'la restricción lado = ancho = alto.', italic=True, color='555555')

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save('DiagramasDeClase2_Solucion.docx')
print('✓ Archivo generado: DiagramasDeClase2_Solucion.docx')